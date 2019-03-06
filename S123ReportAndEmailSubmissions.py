import os
import smtplib
from email import encoders
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import urllib
import requests
import json
import sys
import docx
import glob
import datetime
from arcgis.gis import GIS
from arcgis.apps.survey123._survey import SurveyManager, Survey

# --------------------------------------------------------------------------------------------------------------------------
# 
# This script takes an input Survey ID (from Survey123), a report template (uploaded to the Survey Reports [Beta] tab), and will do the following:
#
# - Generate .docx reports for any new Survey submissions from the last 24 hours,
# - Save these reports in bulk to AGOL, download them one by one to a location, then remove them from AGOL when finished,
# - Read and extract an email address from each .docx file, 
# - Send the relevant .docx file as an attachment to the relevant recipient, 
# - Remove the .docx file once the email has sent,
# - Logs the daily results to a txt file in the output folder.
#
# Developed by John Stowell, 2019
#
# Call this script with     python "..\S123ReportAndEmailSubmissions.py"     
# - eg. in a Unix Cron job or Windows Task Scheduler that runs once a day - since we are always looking for submissions from the last 24 hours
#
#
# Customise these variables in the code below: 
# username, password, attachmentsGeoJSON, and outFolder
#
#
# Also, this script generates the KeyError: 'results' but still works due to the try/except/finally block...
# Related to this ESRI bug:
# BUG-000119057 : The Python API 1.5.2 generate_report() method of the arcgis.apps.survey123 module, generates the following error: { KeyError: 'results' }
#
# API docs: https://esri.github.io/arcgis-python-api/apidoc/html/arcgis.apps.survey123.html
#
# --------------------------------------------------------------------------------------------------------------------------

def main():
   # Customise the variables below
   # --- AGOL information... ---
   org = 'https://YOUR-ORGANISATION.maps.arcgis.com'
   username = 'ARCGIS ONLINE USERNAME'
   password = 'ARCGIS ONLINE PASSWORD'


   # --- Survey123 variables... ---
   surveyID = 'ID OF SURVEY123 FORM' # ID of desired Survey123 form - a unique ID like 4c1b359c4e294c54a02b262b42413f17
   output_folder = r'C:\GISWORK\_tmp\Reports' # Output folder WITHOUT trailing slash. This is also where the log file is stored.

   # WHERE_FILTER: Use '1=1' to return all records, or something like  {{"where":"<col>='<value>'"}  - supports SQL syntax
   # Docs for date queries: https://www.esri.com/arcgis-blog/products/api-rest/data-management/querying-feature-services-date-time-queries/
   # In our case below, we filter by records created in the last 1 day. This works for us as the script is run on a daily schedule.
   where_filter = '{"where":"CreationDate >= CURRENT_TIMESTAMP - INTERVAL \'1\' DAY"}'

   utc_offset = '+13:00' # UTC Offset for location (+13 is NZST)
   report_title = 'Daily_Export' # Title that will show in S123 recent task list
   report_template = 1 # ID of the print template in Survey123 that you want to use (0 = ESRI's sample, 1 = first custom report, 2 = second custom report, etc)


   # --- Email SMTP settings... ---
   email_user = 'EMAIL ADDRESS' # Eg. user@gmail.com. Requires a valid SMTP-enabled email account (Eg. a Gmail acct with the SMTP settings below)
   email_password = 'EMAIL ACCOUNT PASSWORD' # Password for the email account
   smtp_server = 'smtp.gmail.com'
   smtp_port = 587

   # --------------------------------------------------------------------------------------------------------------------------
   # Don't edit below this line - unless you know what you are doing :)
   # --------------------------------------------------------------------------------------------------------------------------



   log = output_folder+"\daily_export_log.txt"
   # Create a log file if it doesn't exist
   print('', file=open(log, "a+"))
   
   # Date variables for later use
   today = datetime.datetime.today()
   yesterday = today-datetime.timedelta(1)

   # -------------------------------------------------------------
   # REPORT GENERATION AND DOWNLOAD PROCESS

   # Initialise AGOL login by script
   print('--------------------------------------------------------------------------------------------------------------------------', file=open(log, "a"))
   print('--- STARTING REPORT GENERATION PROCESS ---', today, file=open(log, "a"))
   print('')
   print('Initialising session in AGOL', file=open(log, "a"))
   print('')
   agol_login = GIS(org, username, password) 

   print('Reading Survey123 information for ID: ',surveyID, file=open(log, "a"))
   print('')
   surveymgr = SurveyManager(agol_login)
   survey = surveymgr.get(surveyID)
   # print('Templates available: ',survey.report_templates) # Return all available print templates for the survey
   # print('')

   template = survey.report_templates[report_template] 
   print('Selected template: ',template, file=open(log, "a"))
   print('')

   reportCount = 0
   # Try/except/finally block to workaround the KeyError: 'results' bug in the generate_report method
   # (Waiting on ESRI to fix this bug.)
   try:
      print('Generating report(s) for submissions from last 24 hours', file=open(log, "a"))
      print('')
      ## Original Example: survey.generate_report(template, '1=1') #generates report for all features
      ## API Docu Example: survey.generate_report(report_template: arcgis.gis.Item, where: str = '1=1', utc_offset: str = '+00:00', report_title: str = None, folder_id: str = None)
      ## Our Example:      survey.generate_report(template, '1=1', '+13:00', 'Test_Report_Export') 
      survey.generate_report(template, where_filter, utc_offset, report_title) 
   except Exception as e:
      print('>> ERROR: KeyError: ',e,' (related to ESRI BUG-000119057)', file=open(log, "a"))
      print('>> Continuing...', file=open(log, "a"))
      print('')
      pass
   finally:
      print('Downloading relevant report(s) to: ',output_folder, file=open(log, "a"))
      print('')
      # Find all Microsoft Word doc files in AGOL with "Survey 123" in the tags
      for x in survey.reports: 
         # Find the creation date (Unix epoch) and convert to local time
         creationdate = datetime.datetime.fromtimestamp(x.created / 1e3)

         # Only find and download AGOL reports created in the last 24 hours 
         # (this will download reports created manually, as well as ones generated by this script)
         if (creationdate > yesterday): 
            # print('Created epoch ',x.created) # Uncomment datestamps below for testing
            # print('Created converted ',creationdate)
            # print('Today converted ',today)
            # print('Yesterday timedelta ',yesterday)
            reportCount += 1
            # Only return reports that contain the surveyID in the html code of our description
            # This should normally return reports generated with the generate_report() method
            if surveyID in x.description: 
               print('Report desc: ',x.description, file=open(log, "a"))
               print('')
               id = x.id # Get ID of each Word doc AGOL item
               data_item = agol_login.content.get(id) 
               data_item.download(save_path = output_folder) # Download each Word doc to specified location
               data_item.delete() # Delete each Word doc item in AGOL (after download finished/no longer required)
      # Finally block end
      print('REPORTS GENERATED: ',reportCount, file=open(log, "a"))
      print('--- REPORT GENERATION PROCESS - FINISHED ---', file=open(log, "a"))
      print('', file=open(log, "a"))






   # -------------------------------------------------------------
   # EMAIL REPORT TO USERS PROCESS

   # Optional - email the report documents to specified email address.
   # Now we cycle through the new .docx reports in our output_folder, extract the user email address and send the attachment to the email that was collected with S123

   print('--- STARTING EMAIL PROCESS ---', file=open(log, "a"))
   sender = email_user
   documentCount = 0

   print('')
   print('Getting list of Word docx files in: ',output_folder, file=open(log, "a"))
   # Add all files ending with .docx to a new list
   file_list = glob.glob(output_folder+'\*.docx')
   print('Files:', file=open(log, "a"))
   for file_name in file_list:
      documentCount += 1
      print(file_name, file=open(log, "a"))

   print('')
   print('Reading raw table data from Word document(s)', file=open(log, "a"))
   print('')
   for file_name in file_list:
      #print(filename)
      # Here we read the tabular data from within our docx report - this is based on the report template that you have created...
      # In my case, there's a table with 10 rows, and row 8 happens to have the Email address that was collected by Survey123.
      docx_data = readDocxTables(file_name)

      # Data[7] happens to be our "Email" row (row 8) in the table within each docx template
      #print('Data7: ',data[7])
      recipient = [str(v) for k,v in docx_data[7].items()][0]
      print('Sending email with attachment to recipient: ',recipient, file=open(log, "a"))

      # Initialise the email and create the enclosing (outer) message
      outer = MIMEMultipart()
      outer['Subject'] = 'Survey Report Attached ' + str(today)
      outer['To'] = recipient
      outer['From'] = sender
      outer.preamble = 'You will not see this in a MIME-aware mail reader.\n'
      msg_text = 'From ORGANISATION:\n\nPlease find the attached Survey Report from our recent visit or discussion. This is a copy of the information discussed and collected for your records.\n\nKind regards,\nYOUR ORGANISATION (email@org.com)\n\n\n\nNOTE: This is an automated message, please do not reply.'

      # Add the attachment to the message
      try:
         with open(file_name, 'rb') as fp:
            msg = MIMEBase('application', 'octet-stream')
            msg.set_payload(fp.read())
         encoders.encode_base64(msg)
         msg.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file_name))
         outer.attach(msg)
         outer.attach(MIMEText(msg_text, 'plain')) # or 'html'
      except:
         print('Unable to open one of the attachments. Error: ', sys.exc_info()[0], file=open(log, "a"))
         raise

      composed = outer.as_string()

      # Send the email via SMTP - we're using Google SMTP servers below
      try:
         with smtplib.SMTP(smtp_server, smtp_port) as s:
               s.ehlo()
               s.starttls()
               s.ehlo()
               s.login(sender, email_password)
               s.sendmail(sender, recipient, composed)
               s.close()
         # Email sent, now let's remove the file so that it isn't sent again the following day
         os.remove(file_name)
         print('Email sent to recipient and removed file from download location.', file=open(log, "a"))
         print('')
      except:
         print('Unable to send the email. Error: ', sys.exc_info()[0], file=open(log, "a"))
         raise

   print('DOCUMENTS SENT TO RECIPIENTS: ',documentCount, file=open(log, "a"))
   print('--- EMAIL PROCESS - FINISHED --- ', file=open(log, "a"))
   print('', file=open(log, "a"))
   print('', file=open(log, "a"))





# -------------------------------------------------------------
# OTHER FUNCTIONS

# Function to retrieve raw text only
def readText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

## Function to retrieve text from tables
def readDocxTables(filename):
    document = docx.Document(filename)
    table = document.tables[1] # 0 = logo, 1 = first block

    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return(data)


if __name__ == '__main__':
    main()
